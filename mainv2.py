# Common imports
import sqlite3
import pysqlite3
from docx import Document
import numpy as np
import PyPDF2
from PyPDF2 import PdfReader
import sys
sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
import os
import openai
from openai import OpenAI
from dotenv import load_dotenv
import json
import lolviz
import requests
# Import the key CrewAI classes
import streamlit as st
import tiktoken
from helper_functions.utility import check_password  

#set up for langchain
import os
from langchain_openai import ChatOpenAI
from langchain_openai import OpenAIEmbeddings



#setup for nativeRag
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader
from crewai_tools import WebsiteSearchTool
from crewai import Agent, Task, Crew
# embedding model that we will use for the session
embeddings_model = OpenAIEmbeddings(model='text-embedding-3-small')

from crewai_tools import (FileReadTool)
# file_tool=FileReadTool()

# llm to be used in RAG pipeplines in this notebook
llm = ChatOpenAI(model='gpt-4o-mini', temperature=0, seed=42)



if load_dotenv('.env'):
   # for local development
   OPENAI_KEY = os.getenv('OPENAI_API_KEY')
else:
   OPENAI_KEY = st.secrets['OPENAI_API_KEY']

client = OpenAI(api_key=OPENAI_KEY)


#function for getting embeddings
def get_embedding(input, model='text-embedding-3-small'):
    response = client.embeddings.create(
        input=input,
        model=model
    )
    return [x.embedding for x in response.data]

#function for getting response

# This is the "Updated" helper function for calling LLM
def get_completion(prompt, model="gpt-4o-mini", temperature=0, top_p=1.0, max_tokens=256, n=1, json_output=False):
    if json_output == True:
      output_json_structure = {"type": "json_object"}
    else:
      output_json_structure = None

    messages = [{"role": "user", "content": prompt}]
    response = client.chat.completions.create( #originally was openai.chat.completions
        model=model,
        messages=messages,
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
        n=1,
        response_format=output_json_structure,
    )
    return response.choices[0].message.content


#get completion by messages

# This a "modified" helper function that we will discuss in this session
# Note that this function directly take in "messages" as the parameter.
def get_completion_by_messages(messages, model="gpt-4o-mini", temperature=0, top_p=1.0, max_tokens=1024, n=1):
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        top_p=top_p,
        max_tokens=max_tokens,
        n=1
    )
    return response.choices[0].message.content

# This function is for calculating the tokens given the "message"
# ⚠️ This is simplified implementation that is good enough for a rough estimation

def count_tokens(text):
    encoding = tiktoken.encoding_for_model('gpt-4o-mini')
    return len(encoding.encode(text))

def count_tokens_from_message_rough(messages):
    encoding = tiktoken.encoding_for_model('gpt-4o-mini')
    value = ' '.join([x.get('content') for x in messages])
    return len(encoding.encode(value))




#Streamlit

# Check if the password is correct.  
if not check_password():  
    st.stop()

#Streamlit section
# region <--------- Streamlit App Configuration --------->
st.set_page_config(
    layout="centered",
    page_title="Understanding the Closure of CPF Special Account"
)
resume=st.file_uploader('''Read me.
                        
                        This app will take in your resume, and a job posting link.
                        once it has both, it will take about 1 minute to generate a customised resume downloadable in word format.
                        you will see the words "Download Analysis as Word Document" Once complete"''')
#define CV & joblink to avoid error so that there is no error

cv_content=None
joblink=None
# Check if a file was uploaded
if resume is not None:
    # Display the file details
    st.write("Uploaded file:", resume.name)

    # Use PdfReader to read the file content
    pdf_reader = PdfReader(resume)
    cv_content = ""

    # Extract text from each page
    for page in pdf_reader.pages:
        cv_content += page.extract_text()

    # Display the extracted content
    st.write("Resume Updated in correct Format")
    
else:
    st.write("Please upload your resume in PDF format.")




joblink=st.text_input('input job posting URL', placeholder="https://www.mycareersfuture.gov.sg/job/information-technology/data-scientist-cpo-03dba75ab1fec49a3aac63d2c676949a?source=MCF&event=Search" )



# Create the agents here
profiler = Agent(
    role="Personal Profiler",

    goal="Conduct comprehensive research on job applicants to help them stand out in the job market.",

    backstory="""The candidates resume is here.{cv_content}
    As a Job Researcher, your expertise in navigating and extracting critical information 
    from job postings is unparalleled. Your skills help identify the necessary qualifications and skills
    sought by employers, forming the foundation for effective application tailoring..""",

    allow_delegation=False, # we will explain more about this later
    # tools=[file_tool],
	verbose=True, # to allow the agent to print out the steps it is taking
)

analyst = Agent(
    role="Tech Job Researcher",

    goal="Perform thorough analysis on job postings to assist job applicants.",

    backstory="""
    Equipped with analytical prowess, you dissect and synthesize information from diverse sources to craft 
    comprehensive personal and professional profiles, laying the groundwork for personalized resume enhancements...""",

    allow_delegation=False, # we will explain more about this later

	verbose=True, # to allow the agent to print out the steps it is taking
)

resume_strategist = Agent(
    role="Resume Strategist",

    goal="Discover the best ways to make a resume stand out in the job market.",

    backstory="""With a strategic mind and an eye for detail, you excel at refining resumes to highlight the most relevant 
    skills and experiences, ensuring they resonate perfectly with the job's requirements.""",

    allow_delegation=False, # we will explain more about this later

	verbose=True, # to allow the agent to print out the steps it is taking
)
# You can use this cell to create tools or functions that you will use in the main code
# Hint: You may want to use `FileReadTool` tool which can be used by agent/task to read the resume file

extract_requirements = Task(
    description="""\
    Analyze the job posting URL provided (`{joblink}`) to extract key skills, experiences, 
    and qualifications required. Use the tools to gather content and identify and categorize the requirements.""",

    expected_output="""\
    A structured list of job requirements, including necessary skills, qualifications, and experiences.""",

    agent=analyst,
)

compile_profile = Task(
    description="""\
    Compile a detailed personal and professional profile based on the current CV.""",

    expected_output="""\
    A comprehensive profile document that includes skills, project experiences, contributions, interests, and communication style.""",

    agent=profiler,
)

align_with_requirement = Task(
    description="""\
    Using the profile and job requirements obtained from previous tasks, tailor the resume to 
    highlight the most relevant areas. Employ tools to adjust and enhance the resume content. Make sure 
    this is the best resume ever but don't make up any information. Update every section, including the initial summary, 
    work experience, skills, and education. All to better reflect the candidate's abilities and how it matches the job posting.""",

    expected_output="""\
    An updated resume that effectively highlights the candidate's qualifications and experiences relevant to the job.""",

    agent=resume_strategist,
)
crew = Crew(
    agents=[profiler,analyst, resume_strategist],
    tasks=[extract_requirements, compile_profile,align_with_requirement],
    verbose=True
)



### this execution will take a few minutes to run
# 
def save_to_word(content, filename="Job_Application_Analysis.docx"):
    doc = Document()
    doc.add_heading("Job Application Analysis", 0)
    
    # Add content (assuming content is a string)
    doc.add_paragraph(str(content))
    
    # Save the document
    doc.save(filename)
    
    return filename

# After running crew.kickoff
if cv_content and joblink:
    job_application_inputs = {
        'joblink': joblink,
        'cv_content': cv_content
    }
    result = crew.kickoff(inputs=job_application_inputs)
    
    # Display the result in Streamlit
    # st.write(result)

    # Save the result as a Word file
    word_file = save_to_word(result, "Job_Application_Analysis.docx")
    
     # Provide a download link for the Word file
    with open(word_file, "rb") as file:
        st.download_button(
            label="Download Analysis as Word Document",
            data=file,
            file_name="Job_Application_Analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.write("Please upload a resume and enter the job link.")